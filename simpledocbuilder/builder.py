"""
Module: builder.py
Description: Cœur du système SimpleDocBuilder.
Fournit une classe principale `SimpleDocBuilder` pour générer des documents Word
de manière modulaire et robuste.
"""

import logging
import tempfile
import sys
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Union, Any, Dict, Iterator

import pandas as pd
from docx import Document
from docx.document import Document as _Document
from docx.shared import Mm
from docxcompose.composer import Composer
from docxtpl import DocxTemplate, InlineImage
from PIL import Image as PILImage

# --- Gestion des dépendances optionnelles ---
try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

try:
    import imgkit
    IMGKIT_AVAILABLE = True
except ImportError:
    IMGKIT_AVAILABLE = False

# Import des utilitaires locaux
try:
    from simpledocbuilder.utils import eng_string, auto_crop_simple, random_name
except ImportError:
    # Cas où simpledocbuilder est dans le path courant
    from utils import eng_string, auto_crop_simple, random_name

# Configuration Logging
logger = logging.getLogger("simpledocbuilder")
if not logger.handlers:
    # Configuration par défaut si non configuré ailleurs
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.INFO)


# --- Exceptions et Config ---

class DocumentError(Exception):
    """Exception de base pour les erreurs de génération de document."""
    pass


@dataclass
class DocumentConfig:
    """Configuration globale pour le SimpleDocBuilder."""
    temp_prefix: str = 'sdb_'
    default_image_width_mm: int = 150
    default_table_header: str = ''
    default_index_name: str = ''
    logging_level: int = logging.INFO


# --- Abstraction Block ---

class ContentBlock(ABC):
    """
    Classe abstraite représentant un bloc de contenu (Texte, Image, Tableau, etc.).
    Chaque bloc est responsable de se rendre lui-même dans un fichier .docx temporaire.
    """
    @abstractmethod
    def render(self, output_path: Path, context: Dict[str, Any] = None) -> Path:
        """
        Génère le contenu du bloc dans le fichier spécifié par output_path.
        
        Args:
            output_path: Chemin absolu vers le fichier .docx de sortie.
            context: Contexte global optionnel (variables partagées).
            
        Returns:
            Le chemin vers le fichier généré (output_path).
        """
        pass


# --- Implémentations des Blocs ---

@dataclass
class TextBlock(ContentBlock):
    text: str
    style: Optional[str] = None  # ex: 'Heading 1', 'Normal'

    def render(self, output_path: Path, context: Dict = None) -> Path:
        doc = Document()
        if self.text:
            p = doc.add_paragraph(str(self.text))
            if self.style:
                try:
                    p.style = self.style
                except KeyError:
                    logger.warning(f"Style '{self.style}' introuvable, utilisation du style par défaut.")
        doc.save(output_path)
        return output_path


@dataclass
class ImageBlock(ContentBlock):
    path: Union[str, Path]
    width_mm: int = 150
    caption: Optional[str] = None
    template_path: Optional[Union[str, Path]] = None

    def render(self, output_path: Path, context: Dict = None) -> Path:
        image_path = Path(self.path).resolve()
        
        if not image_path.exists():
            # Gestion d'erreur locale : on crée un placeholder pour ne pas planter tout le build
            logger.error(f"Image introuvable : {image_path}")
            doc = Document()
            doc.add_paragraph(f"[ERREUR: Image introuvable - {image_path.name}]", style="Normal")
            doc.save(output_path)
            return output_path

        if self.template_path:
            # Mode Template : injection dans un conteneur existant
            tpl_path = Path(self.template_path).resolve()
            if not tpl_path.exists():
                raise DocumentError(f"Template d'image introuvable : {tpl_path}")
            
            doc = DocxTemplate(tpl_path)
            # On assume que le template attend 'image' et 'title'
            render_context = {
                'image': InlineImage(doc, str(image_path), width=Mm(self.width_mm)),
                'title': self.caption or ""
            }
            # Merge avec le contexte global si nécessaire
            if context:
                render_context = {**context, **render_context}
                
            doc.render(render_context)
            doc.save(output_path)
        else:
            # Mode Simple : ajout direct
            doc = Document()
            if self.caption:
                # Idéalement 'Caption', mais fallback sur Normal si pas présent
                try:
                    doc.add_paragraph(self.caption, style='Caption')
                except:
                    doc.add_paragraph(self.caption)
            try:
                doc.add_picture(str(image_path), width=Mm(self.width_mm))
            except Exception as e:
                logger.error(f"Erreur lors de l'ajout de l'image {image_path}: {e}")
                doc.add_paragraph(f"[ERREUR IMAGE: {e}]")
            doc.save(output_path)
            
        return output_path


@dataclass
class DataFrameBlock(ContentBlock):
    df: pd.DataFrame
    title: Optional[str] = None
    header_col: str = ''
    index_name: str = ''
    formatting: str = '%.2f' # Format ingénieur par défaut
    use_eng_format: bool = True
    template_path: Optional[Union[str, Path]] = None

    def render(self, output_path: Path, context: Dict = None) -> Path:
        if self.df.empty:
            logger.warning("DataFrame vide, génération d'un bloc vide.")
            doc = Document()
            doc.add_paragraph("[Tableau Vide]")
            doc.save(output_path)
            return output_path

        local_df = self.df.copy()

        # Formatage des données numériques
        if self.use_eng_format:
            # Compatibilité Pandas map vs applymap
            method = getattr(local_df, "map", getattr(local_df, "applymap"))
            # Applique eng_string uniquement sur les numériques (int/float)
            local_df = method(lambda x: eng_string(x, formating=self.formatting) if isinstance(x, (int, float)) else x)

        if self.template_path:
            # Mode Template
            doc = DocxTemplate(self.template_path)
            # Préparation structure de données typique pour iterer dans jinja2
            # Structure proposée: table.col_labels, table.tbl_contents -> [{label: index, cols: [vals]}]
            table_data = {
                'col_labels': local_df.columns.tolist(),
                'tbl_contents': [
                    {'label': idx, 'cols': row.tolist()} 
                    for idx, row in local_df.iterrows()
                ]
            }
            render_context = {
                'table': table_data,
                'title': self.title or "",
                'header_col': self.header_col,
                'name_index': self.index_name
            }
            if context:
                render_context = {**context, **render_context}
            
            doc.render(render_context)
            doc.save(output_path)
            
        else:
            # Mode Natif (Word Table)
            doc = Document()
            if self.title:
                doc.add_heading(self.title, level=2)
            
            # +1 pour l'index
            rows_count = len(local_df) + 1
            cols_count = len(local_df.columns) + 1
            
            table = doc.add_table(rows=rows_count, cols=cols_count)
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = str(self.index_name)
            for i, col in enumerate(local_df.columns):
                hdr_cells[i+1].text = str(col)
            
            # Body
            for i, (idx, row) in enumerate(local_df.iterrows()):
                row_cells = table.rows[i+1].cells
                row_cells[0].text = str(idx)
                for j, val in enumerate(row):
                    row_cells[j+1].text = str(val) if pd.notna(val) else ""
            
            doc.save(output_path)

        return output_path


@dataclass
class TemplateBlock(ContentBlock):
    """Pour injecter des données arbitraires dans un template Word."""
    template_path: Union[str, Path]
    data_context: Dict[str, Any]

    def render(self, output_path: Path, context: Dict = None) -> Path:
        doc = DocxTemplate(self.template_path)
        # Fusion contexte global + contexte local
        full_context = {**(context or {}), **self.data_context}
        doc.render(full_context)
        doc.save(output_path)
        return output_path


@dataclass
class LatexBlock(ContentBlock):
    latex_content: str
    
    def render(self, output_path: Path, context: Dict = None) -> Path:
        if not PYPANDOC_AVAILABLE:
            raise DocumentError("pypandoc requis mais non installé.")
        
        if not self.latex_content.strip():
            # Pas de contenu, on crée un doc vide
            doc = Document()
            doc.save(output_path)
            return output_path

        try:
            pypandoc.convert_text(
                self.latex_content,
                format='latex',
                to='docx',
                outputfile=str(output_path)
            )
        except Exception as e:
            raise DocumentError(f"Erreur conversion LaTeX: {e}")
        return output_path


@dataclass
class HtmlBlock(ContentBlock):
    html_content: str
    
    def render(self, output_path: Path, context: Dict = None) -> Path:
        if not PYPANDOC_AVAILABLE:
            raise DocumentError("pypandoc requis mais non installé.")
        
        if not self.html_content.strip():
            doc = Document()
            doc.save(output_path)
            return output_path
            
        try:
            pypandoc.convert_text(
                self.html_content,
                format='html',
                to='docx',
                outputfile=str(output_path)
            )
        except Exception as e:
            raise DocumentError(f"Erreur conversion HTML: {e}")
        return output_path


@dataclass
class GreatTableBlock(ContentBlock):
    """Gère le rendu de HTML (issu de Great Tables ou autre) via conversion Image (imgkit)."""
    html_str: str
    width_mm: int = 150
    title: Optional[str] = None
    template_path: Optional[Union[str, Path]] = None
    temp_dir_gen: Any = None # Nécessaire pour générer l'image temp dans le même dossier

    def render(self, output_path: Path, context: Dict = None) -> Path:
        if not IMGKIT_AVAILABLE:
            raise DocumentError("imgkit requis pour GreatTableBlock.")
        
        # 1. Générer l'image depuis le HTML
        # On a besoin d'un chemin temporaire pour l'image
        # On utilise le même dossier que output_path pour être propre
        img_filename = output_path.with_suffix('.png')
        
        try:
            imgkit.from_string(self.html_str, str(img_filename), options={'quiet': ''})
        except Exception as e:
            logger.error(f"Echec imgkit: {e}")
            # Tente de continuer ou fallback
            doc = Document()
            doc.add_paragraph(f"[ERREUR IMGKIT: {e}]")
            doc.save(output_path)
            return output_path

        # 2. Rogner l'image (auto-crop)
        cropped = auto_crop_simple(img_filename)
        if cropped:
            cropped.save(img_filename)
        
        # 3. Insérer l'image (réutilisation de logic ImageBlock simplifiée)
        # On délègue à ImageBlock pour ne pas dupliquer la logique template/simple
        img_block = ImageBlock(
            path=img_filename, 
            width_mm=self.width_mm, 
            caption=self.title, 
            template_path=self.template_path
        )
        return img_block.render(output_path, context)


# --- Le Builder Principal ---

class SimpleDocBuilder:
    """
    Classe principale pour la construction de documents Word.
    S'utilise de préférence avec un block 'with'.
    """

    def __init__(self, config: Optional[DocumentConfig] = None):
        self.config = config or DocumentConfig()
        self.blocks: List[ContentBlock] = []
        self._temp_dir: Optional[tempfile.TemporaryDirectory] = None
        self._name_gen = random_name()
        self.global_context: Dict[str, Any] = {}
        
        # Validation du logging
        logger.setLevel(self.config.logging_level)

    def init(self):
        """Initialisation explicite du dossier temporaire."""
        if self._temp_dir is None:
            self._temp_dir = tempfile.TemporaryDirectory(prefix=self.config.temp_prefix)
            logger.info(f"Dossier temporaire initialisé: {self._temp_dir.name}")
        return self

    def cleanup(self):
        """Nettoyage des ressources."""
        if self._temp_dir:
            try:
                self._temp_dir.cleanup()
                logger.info("Dossier temporaire nettoyé.")
            except Exception as e:
                logger.warning(f"Erreur lors du nettoyage du dossier temporaire: {e}")
            self._temp_dir = None

    def __enter__(self):
        self.init()
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.cleanup()

    def _get_next_path(self) -> Path:
        """Retourne un chemin unique pour le prochain fragment."""
        if not self._temp_dir:
            self.init() # Auto-init si oublié
        return Path(self._temp_dir.name) / f"{next(self._name_gen)}.docx"

    # --- API Fluent ---

    def add(self, block: ContentBlock) -> 'SimpleDocBuilder':
        """Ajoute un bloc générique à la liste de construction."""
        self.blocks.append(block)
        return self

    def add_text(self, text: str, style: str = None) -> 'SimpleDocBuilder':
        """Ajoute un paragraphe de texte."""
        return self.add(TextBlock(text, style))
    
    def add_title(self, text: str, level: int = 1) -> 'SimpleDocBuilder':
        """Ajoute un titre (Heading 1, Heading 2, etc.)."""
        style = f'Heading {level}'
        return self.add(TextBlock(text, style))

    def add_image(self, path: Union[str, Path], width_mm: Optional[int] = None, 
                 title: str = None, template_path: str = None) -> 'SimpleDocBuilder':
        """Ajoute une image."""
        w = width_mm if width_mm is not None else self.config.default_image_width_mm
        return self.add(ImageBlock(path, w, title, template_path))

    def add_table(self, df: pd.DataFrame, title: str = None, use_eng_format: bool = True, 
                  template_path: str = None) -> 'SimpleDocBuilder':
        """Ajoute un tableau Pandas."""
        return self.add(DataFrameBlock(
            df=df, 
            title=title, 
            header_col=self.config.default_table_header,
            index_name=self.config.default_index_name,
            use_eng_format=use_eng_format,
            template_path=template_path
        ))

    def add_latex(self, latex: str) -> 'SimpleDocBuilder':
        """Ajoute du contenu LaTeX (via pandoc)."""
        return self.add(LatexBlock(latex))

    def add_html(self, html: str) -> 'SimpleDocBuilder':
        """Ajoute du contenu HTML (via pandoc, rendu natif)."""
        return self.add(HtmlBlock(html))

    def add_gt_html(self, html: str, title: str = None, width_mm: Optional[int] = None, 
                   template_path: str = None) -> 'SimpleDocBuilder':
        """Ajoute du HTML "complexe" (ex: Great Tables) rendu comme image."""
        w = width_mm if width_mm is not None else self.config.default_image_width_mm
        return self.add(GreatTableBlock(html, w, title, template_path))

    def add_template(self, template_path: str, context: Dict[str, Any]) -> 'SimpleDocBuilder':
        """Ajoute un template brut avec des données."""
        return self.add(TemplateBlock(template_path, context))
    
    # --- Construction ---

    def build(self, output_path: Union[str, Path]) -> None:
        """
        Génère tous les fragments et les assemble dans le fichier final.
        """
        if not self.blocks:
            logger.warning("Aucun contenu à générer. Fichier de sortie non créé.")
            return

        final_path = Path(output_path).resolve()
        
        # S'assurer que le dossier de sortie existe
        if not final_path.parent.exists():
            final_path.parent.mkdir(parents=True, exist_ok=True)

        logger.info(f"Début de la génération de {len(self.blocks)} blocs.")
        
        master_doc = Document()
        composer = Composer(master_doc)
        
        # Indicateur pour savoir si c'est le premier document ajouté
        # (docxcompose conseille d'append sur un master, mais si master vide, 
        # le premier doc définit parfois les styles de base mieux que le vide)
        is_first = True

        for i, block in enumerate(self.blocks):
            try:
                # 1. Rendu individuel
                fragment_path = self._get_next_path()
                block.render(fragment_path, context=self.global_context)
                
                # 2. Composition
                # On utilise append pour tout le monde. Composer gère les sauts de section.
                # Pour plus de robustesse, on re-ouvre le docx généré pour le valider
                if fragment_path.exists():
                    try:
                        part_doc = Document(str(fragment_path))
                        composer.append(part_doc)
                        logger.debug(f"Bloc {i+1}/{len(self.blocks)} ajouté ({type(block).__name__}).")
                    except Exception as e:
                        logger.error(f"Erreur lors de l'assemblage du bloc {i+1}: {e}")
                else:
                    logger.error(f"Erreur: Le fichier fragment {fragment_path} n'a pas été généré.")

            except Exception as e:
                logger.error(f"Echec critique sur le bloc {i}: {e}")
                # On continue pour essayer de sauver ce qui peut l'être ?
                # Ou on raise ? Spec dit "build fusionne". Si un bloc fail, le doc est incomplet.
                # On choisit de continuer en logguant l'erreur pour la robustesse (partial result > no result)
                continue

        try:
            composer.save(str(final_path))
            logger.info(f"Document final généré avec succès : {final_path}")
        except Exception as e:
            raise DocumentError(f"Erreur lors de la sauvegarde finale : {e}")
